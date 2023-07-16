
import docx
import re
import sys
import tika
from tika import parser
tika.initVM()
from datetime import datetime
import ycalc
from skill_parser import ResumeParser
def expgen(resume_file):
    final1=[]
    final2=[]
    final3=[]
    split=0
    a=resume_file
    p1=[]
    p2=[]
    p3=[]
    sy1=[]
    sm1=[]
    sy2=[]
    sm2=[]
    sy3=[]
    sm3=[]
    exp1=[]
    exp2=[]
    exp3=[]
    pop1=[]
    pop2=[]
    pop3=[]
    c1=-1
    c2=-1
    c3=-1
    se_sc=[]
    
    search=["transformer 7.0","html","sap bw","Cognos 11"]
    def concatenate_list_data(list):    #to convert list into one string
        result= ''
        for element in list:
            result += str(element)
        return result
    def same_start_date(m,y,mlist,ylist):             #not required
        if(mlist.count(m)==1 and ylist.count(y)==1):
            return 1
        else:
            return 0
    if(a.endswith(".docx")):                   # checking for .docx pattern
        def getTextTables(filename):
            document = docx.Document(a)
            tables = document.tables
            fullText=[]
            if tables:
                
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                fullText.append(paragraph.text)
                return '\n'.join(fullText)
        intermediate1=getTextTables(a)
        


        if (intermediate1):
            ans1=intermediate1.splitlines()
            
            
            
            for line in ans1:
                c1=c1+1
                
                x=re.findall(r"[0-9]*\.*[0-9]\s*\+*\s*(?i)year[s]?",line)
                if x:
                    print(x)
            

                for i in x:
                    c=0
                    while(i[c].isdigit()):
                        c=c+1
                    final1.append(i[0:c])
                
                


                x=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)[ ]*.[ ]*[0-9]{2,4}",line)
                
                if(x):
                    print(x)
                    print("Found at index position",c1)
                    p1.append(c1)
                    

                    if(len(x)==1):
                        firstyear=re.findall(r"[0-9]{2,4}",x[0])
                        firstmonth=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)",x[0])
                        lsecondmonth=datetime.now().strftime('%h')
                        secondmonth=[]
                        secondmonth.append(lsecondmonth)
                        lsecondyear=str(datetime.now().year)
                        secondyear=[]
                        secondyear.append(lsecondyear)
                        print(firstmonth)
                        print(firstyear)
                        print(secondmonth)
                        print(secondyear)
                        exp=ycalc.expcalc1(firstmonth[0],firstyear[0],secondmonth[0],secondyear[0])
                        if(exp>0 and exp<35):
                            exp1.append(exp)
                            sy1.append(firstyear)
                            sm1.append(firstmonth)
                        else:
                            p1.pop()    
                            

                    else:
                            
                        firstyear=re.findall(r"[0-9]{2,4}",x[0])
                        firstmonth=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)",x[0])
                        secondyear=re.findall(r"[0-9]{2,4}",x[1])
                        secondmonth=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)",x[1])
                        print(firstmonth)
                        print(firstyear)
                        print(secondmonth)
                        print(secondyear)
                        exp=ycalc.expcalc1(firstmonth[0],firstyear[0],secondmonth[0],secondyear[0])
                        if(exp>0 and exp<35):
                            exp1.append(exp)
                            sy1.append(firstyear)
                            sm1.append(firstmonth)
                        else:
                            p1.pop() 
                        

                x=re.findall(r"\s+[01]?[0-9]{1}[-/]1?9?2?0?[0-9]{2}",line)
                
                if(x):
                    print(x)
                    print("Found at index position",c1)
                    p1.append(c1)
                    

                    if(len(x)==1):
                        d1=re.split('/|-', x[0])
                        firstmonth=d1[0]
                        firstyear=d1[1]
                        secondmonth=datetime.now().month
                        
                        
                        secondyear=datetime.now().year
                        
                        
                        print(firstmonth)
                        print(firstyear)
                        print(secondmonth)
                        print(secondyear)
                        exp=ycalc.expcalc2(firstmonth,firstyear,secondmonth,secondyear)
                        if(exp>0 and exp<35):
                            exp1.append(exp)
                            sy1.append(firstyear)
                            sm1.append(firstmonth)
                        else:
                            p1.pop() 
                        
                            

                    else:
                        d1=re.split('/|-', x[0])
                        firstmonth=d1[0]
                        firstyear=d1[1]
                        d2=re.split('/|-', x[1])
                        secondmonth=d2[0]
                        secondyear=d2[1]
                        print(firstmonth)
                        print(firstyear)
                        print(secondmonth)
                        print(secondyear)
                        exp=ycalc.expcalc2(firstmonth,firstyear,secondmonth,secondyear)
                        if(exp>0 and exp<35):
                            exp1.append(exp)
                            sy1.append(firstyear)
                            sm1.append(firstmonth)
                        else:
                            p1.pop() 
                              
            
          
        def getText(filename):
            doc = docx.Document(filename)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            return '\n'.join(fullText)


        intermediate2=getText(a)
        ans2=intermediate2.splitlines()
        
        
        
        
        for line in ans2:
            c2=c2+1
            x=re.findall(r"[0-9]*\.*[0-9]\s*\+*\s*(?i)year[s]?",line)
            if x:
                print(x)
            

            for i in x:
                c=0
                while(i[c].isdigit()):
                    c=c+1
                final2.append(i[0:c])
            


            x=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)[ ]*.[ ]*[0-9]{2,4}",line)
                    
            if(x):

                print(x)
                print("Found at index position",c2)
                p2.append(c2)
                if(len(x)==1):
                    firstyear=re.findall(r"[0-9]{2,4}",x[0])
                    firstmonth=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)",x[0])
                    lsecondmonth=datetime.now().strftime('%h')
                    secondmonth=[]
                    secondmonth.append(lsecondmonth)
                    lsecondyear=str(datetime.now().year)
                    secondyear=[]
                    secondyear.append(lsecondyear)
                    print(firstmonth)
                    print(firstyear)
                    print(secondmonth)
                    print(secondyear)
                    exp=ycalc.expcalc1(firstmonth[0],firstyear[0],secondmonth[0],secondyear[0])
                    if(exp>0 and exp<35):
                        exp2.append(exp)
                        sy2.append(firstyear)
                        sm2.append(firstmonth)
                    else:
                        p2.pop() 
                    
                            

                else:
                            
                    firstyear=re.findall(r"[0-9]{2,4}",x[0])
                    firstmonth=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)",x[0])
                    secondyear=re.findall(r"[0-9]{2,4}",x[1])
                    secondmonth=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)",x[1])
                    print(firstmonth)
                    print(firstyear)
                    print(secondmonth)
                    print(secondyear)
                    exp=ycalc.expcalc1(firstmonth[0],firstyear[0],secondmonth[0],secondyear[0])
                    if(exp>0 and exp<35):
                        exp2.append(exp)
                        sy2.append(firstyear)
                        sm2.append(firstmonth)
                    else:
                        p2.pop() 
                    



            x=re.findall(r"\s+[01]?[0-9]{1}[-/]1?9?2?0?[0-9]{2}",line)
            if(x):
                print(x)
                print("Found at index position",c2)
                p2.append(c2)
                    

                if(len(x)==1):
                    d1=re.split('/|-', x[0])
                    firstmonth=d1[0]
                    firstyear=d1[1]
                    secondmonth=datetime.now().month
                        
                        
                    secondyear=datetime.now().year
                        
                        
                    print(firstmonth)
                    print(firstyear)
                    print(secondmonth)
                    print(secondyear)
                    exp=ycalc.expcalc2(firstmonth,firstyear,secondmonth,secondyear)
                    if(exp>0 and exp<35):
                        exp2.append(exp)
                        sy2.append(firstyear)
                        sm2.append(firstmonth)
                    else:
                        p2.pop() 
                    
                            

                else:
                    d1=re.split('/|-', x[0])
                    firstmonth=d1[0]
                    firstyear=d1[1]
                    d2=re.split('/|-', x[1])
                    secondmonth=d2[0]
                    secondyear=d2[1]
                    print(firstmonth)
                    print(firstyear)
                    print(secondmonth)
                    print(secondyear)
                    exp=ycalc.expcalc2(firstmonth,firstyear,secondmonth,secondyear)
                    if(exp>0 and exp<35):
                        exp2.append(exp)
                        sy2.append(firstyear)
                        sm2.append(firstmonth)
                    else:
                        p2.pop() 
                      
        
                
    if(a.endswith(".pdf")):
        raw = parser.from_file(a)
        intermediate3=raw['content']
        ans3=intermediate3.splitlines()
        
        
        for line in ans3:
            c3=c3+1
            x=re.findall(r"[0-9]*\.*[0-9]\s*\+*\s*(?i)year[s]?",line)
            if x:
                print(x)
            

            for i in x:
                c=0
                while(i[c].isdigit()):
                    c=c+1
                final3.append(i[0:c])
            

            x=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)[ ]*.[ ]*[0-9]{2,4}",line)
                    
            if(x):
                print(x)
                print("Found at index position",c3)
                p3.append(c3)
                if(len(x)==1):
                    firstyear=re.findall(r"[0-9]{2,4}",x[0])
                    firstmonth=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)",x[0])
                    lsecondmonth=datetime.now().strftime('%h')
                    secondmonth=[]
                    secondmonth.append(lsecondmonth)
                    lsecondyear=str(datetime.now().year)
                    secondyear=[]
                    secondyear.append(lsecondyear)
                    print(firstmonth)
                    print(firstyear)
                    print(secondmonth)
                    print(secondyear)
                    exp=ycalc.expcalc1(firstmonth[0],firstyear[0],secondmonth[0],secondyear[0])
                    if(exp>0 and exp<35):
                        exp3.append(exp)
                        sy3.append(firstyear)
                        sm3.append(firstmonth)
                    else:
                        p3.pop() 
                    
                            

                else:
                            
                    firstyear=re.findall(r"[0-9]{2,4}",x[0])
                    firstmonth=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)",x[0])
                    secondyear=re.findall(r"[0-9]{2,4}",x[1])
                    secondmonth=re.findall(r"\b(?i)(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|(?:Nov|Dec)(?:ember)?)",x[1])
                    print(firstmonth)
                    print(firstyear)
                    print(secondmonth)
                    print(secondyear)
                    exp=ycalc.expcalc1(firstmonth[0],firstyear[0],secondmonth[0],secondyear[0])
                    if(exp>0 and exp<35):
                        exp3.append(exp)
                        sy3.append(firstyear)
                        sm3.append(firstmonth)
                    else:
                        p3.pop() 
                    



            x=re.findall(r"\s+[01]?[0-9]{1}[-/]1?9?2?0?[0-9]{2}",line)
            if(x):
                print(x)
                print("Found at index position",c3)
                p3.append(c3)
                    

                if(len(x)==1):
                    d1=re.split('/|-', x[0])
                    firstmonth=d1[0]
                    firstyear=d1[1]
                    secondmonth=datetime.now().month
                        
                        
                    secondyear=datetime.now().year
                        
                        
                    print(firstmonth)
                    print(firstyear)
                    print(secondmonth)
                    print(secondyear)
                    exp=ycalc.expcalc2(firstmonth,firstyear,secondmonth,secondyear)
                    if(exp>0 and exp<35):
                        exp3.append(exp)
                        sy3.append(firstyear)
                        sm3.append(firstmonth)
                    else:
                        p3.pop() 
                    
                            

                else:
                    d1=re.split('/|-', x[0])
                    firstmonth=d1[0]
                    firstyear=d1[1]
                    d2=re.split('/|-', x[1])
                    secondmonth=d2[0]
                    secondyear=d2[1]
                    print(firstmonth)
                    print(firstyear)
                    print(secondmonth)
                    print(secondyear)
                    exp=ycalc.expcalc2(firstmonth,firstyear,secondmonth,secondyear)
                    if(exp>0 and exp<35):
                        exp3.append(exp)
                        sy3.append(firstyear)
                        sm3.append(firstmonth)
                    else:
                        p3.pop() 
                      
            
    if(not final1):
        final3.append("0")   
    if(not final2):
        final2.append("0")
    if(not final3):
        final3.append("0")
    if p1:
        l=len(p1)
        for x in range(l):
            if(sm1.count(sm1[x])>1 and sy1.count(sy1[x])>1):
                pop1.append(x)
        for index in sorted(pop1, reverse=True):
            del p1[index]
            del exp1[index]
            del sy1[index]
            del sm1[index]

        
    if p2:
       
        l=len(p2)
        print(p2)
        print(exp2)

        for x in range(l):
            
            if(sm2.count(sm2[x])>1 and sy2.count(sy2[x])>1):
                pop2.append(x)
                
                
        for index in sorted(pop2, reverse=True):
            del p2[index]
            del exp2[index]
            del sy2[index]
            del sm2[index]
    if p3:
        l=len(p3)
        for x in range(l):
            if(sm3.count(sm3[x])>1 and sy3.count(sy3[x])>1):
                pop3.append(x)
        for index in sorted(pop3, reverse=True):
            del p3[index]
            del exp3[index]
            del sy3[index]
            del sm3[index]


    print(final1)
    print(final2)
    print(final3)

    m=max(final1,final2,final3)
    print(m)
    print("Total experince considered is:",m[0])
    print(p1)
    print(exp1)
    print(sy1)
    print(sm1)
    print(p2)
    print(exp2)
    print(sy2)
    print(sm2)
    print(p3)
    print(exp3)
    print(sy3)
    print(sm3)
    if p1:
        l=len(p1)
        for x in range(l):

            start=p1[x]
            if(x==l-1):
                print(ans1[start:])
                r=concatenate_list_data(ans1[start:])
                ans=ResumeParser(r,search).get_extracted_data()
                ans["exp"]=exp1[x]
                print(ans)
                if(int(ans["score"]>0)):
                    se_sc.append(ans["score"])
                    se_sc.append(ans["exp"])
                    se_sc.append(m[0])
            else:
                end=p1[x+1]
                print(ans1[start:end])
                r=concatenate_list_data(ans1[start:end])
                ans=ResumeParser(r,search).get_extracted_data()
                ans["exp"]=exp1[x]
                print(ans)
                if(int(ans["score"]>0)):
                    se_sc.append(ans["score"])
                    se_sc.append(ans["exp"])
                    se_sc.append(m[0])
                
    if p2:
        l=len(p2)
        for x in range(l):
            start=p2[x]
            if(x==l-1):
                print(ans2[start:])
                r=concatenate_list_data(ans2[start:])
                ans=ResumeParser(r,search).get_extracted_data()
                ans["exp"]=exp2[x]
                print(ans)
                if(int(ans["score"]>0)):
                    se_sc.append(ans["score"])
                    se_sc.append(ans["exp"])
                    se_sc.append(m[0])
                
            else:
                end=p2[x+1]
                print(ans2[start:end])
                r=concatenate_list_data(ans2[start:end])
                ans=ResumeParser(r,search).get_extracted_data()
                ans["exp"]=exp2[x]
                print(ans)
                if(int(ans["score"]>0)):
                    se_sc.append(ans["score"])
                    se_sc.append(ans["exp"])
                    se_sc.append(m[0])
                
    if p3:
        l=len(p3)
        for x in range(l):
            start=p3[x]
            if(x==l-1):
                print(ans3[start:])
                r=concatenate_list_data(ans3[start:])
                ans=ResumeParser(r,search).get_extracted_data()
                ans["exp"]=exp3[x]
                print(ans)
                if(int(ans["score"]>0)):
                    se_sc.append(ans["score"])
                    se_sc.append(ans["exp"])
                    se_sc.append(m[0])
                
            else:
                end=p3[x+1]
                print(ans3[start:end])
                r=concatenate_list_data(ans3[start:end])
                ans=ResumeParser(r,search).get_extracted_data()
                ans["exp"]=exp3[x]
                print(ans)
                if(int(ans["score"]>0)):
                    se_sc.append(ans["score"])
                    se_sc.append(ans["exp"])
                    se_sc.append(m[0])
    if(len(se_sc)==0):
        se_sc.append(0)
        se_sc.append(0)
        se_sc.append(m[0])
    return se_sc

            




    

